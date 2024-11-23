import streamlit as st
import tempfile
from lxml import etree
import zipfile
from docx import Document
from models import ExtractA, ExtractB, ExtractC, ExtractD
from utils import chunk_by_token


# Lớp CustomDocumentConverter để xử lý file Word
class CustomDocumentConverter:
    def convert(self, file_path):
        """
        Chuyển đổi file .docx thành nội dung văn bản bằng cách giải nén XML.
        """
        try:
            # Đọc nội dung XML từ file .docx
            with zipfile.ZipFile(file_path, 'r') as z:
                xml_content = z.read('word/document.xml')

            # Parse nội dung XML
            root = etree.XML(xml_content)

            # Định nghĩa namespace
            namespaces = root.nsmap
            if "w" not in namespaces:
                namespaces["w"] = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            # Tìm tất cả các phần tử văn bản
            elements = root.xpath(".//w:t", namespaces=namespaces)
            text = " ".join([el.text for el in elements if el.text])

            return text
        except KeyError:
            raise ValueError("File .docx không hợp lệ hoặc không chứa tài liệu Word!")
        except Exception as e:
            raise ValueError(f"Lỗi xử lý file Word: {e}")


# Hàm xuất danh sách câu hỏi ra file Word
def export_to_word(questions):
    """
    Xuất danh sách câu hỏi ra file Word và trả về nội dung file.
    """
    doc = Document()
    doc.add_heading("Danh sách câu hỏi", level=1)

    for idx, quiz in enumerate(questions, start=1):
        doc.add_heading(f"Câu {idx} ({quiz['level']}):", level=2)
        doc.add_paragraph(quiz["question"])
        doc.add_paragraph("Đáp án:")
        for i, choice in enumerate(quiz["choices"], start=1):
            doc.add_paragraph(f"{i}. {choice}")
        doc.add_paragraph(f"Đáp án đúng: {quiz['answer']}")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.seek(0)
    return temp_file


# Giao diện Streamlit
st.title("Trích xuất câu hỏi từ tài liệu Word")

# Tải file lên
uploaded_file = st.file_uploader("Tải lên file .docx", type=["docx"])

if uploaded_file:
    with st.spinner("Đang xử lý tệp..."):
        try:
            # Lưu file tạm
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(uploaded_file.getbuffer())
                temp_file_path = temp_file.name

            # Sử dụng CustomDocumentConverter để xử lý file
            converter = CustomDocumentConverter()
            text = converter.convert(temp_file_path)

            # Chia văn bản thành các chunk
            chunks = chunk_by_token(text, max_len=1000)
            st.success("File đã được xử lý thành công!")

            # Hiển thị các chunk
            st.write("### Chọn chunk để xử lý:")
            selected_chunk = st.slider(
                "Chọn chunk để trích xuất câu hỏi",
                min_value=0,
                max_value=len(chunks) - 1,
                value=0
            )
            st.text_area("Nội dung chunk đã chọn:", chunks[selected_chunk], height=200)

            # Chọn mức độ câu hỏi
            extraction_levels = {
                "Dễ": ExtractA,
                "Thông hiểu": ExtractB,
                "Vận dụng": ExtractC,
                "Vận dụng cao": ExtractD
            }
            selected_levels = st.multiselect(
                "Chọn mức độ câu hỏi cần trích xuất",
                options=list(extraction_levels.keys()),
                default=["Dễ", "Thông hiểu", "Vận dụng", "Vận dụng cao"]
            )

            # Trích xuất câu hỏi
            if st.button("Trích xuất câu hỏi"):
                with st.spinner("Đang trích xuất câu hỏi..."):
                    all_results = []
                    for level in selected_levels:
                        extractor_class = extraction_levels[level]
                        extractor = extractor_class()
                        result = extractor.run(chunks[selected_chunk])
                        all_results.extend(result["quizes"])

                    # Hiển thị kết quả
                    st.write("### Câu hỏi được trích xuất:")
                    for quiz in all_results:
                        st.write(f"#### Câu hỏi ({quiz['level']}):")
                        st.write(quiz["question"])
                        st.write("**Đáp án:**")
                        for i, choice in enumerate(quiz["choices"], 1):
                            st.write(f"{i}. {choice}")
                        st.write(f"**Đáp án đúng:** {quiz['answer']}")

                    # Tải xuống file Word
                    temp_file = export_to_word(all_results)
                    with open(temp_file.name, "rb") as f:
                        st.download_button(
                            label="Tải xuống file Word",
                            data=f,
                            file_name="questions.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

        except Exception as e:
            st.error(f"Lỗi: {e}")
